"""
文档工具模块

支持：
- Markdown 文件的创建、读取、写入、删除
- Markdown 转 Word 文档（保留格式）
"""

import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


class DocumentManager:
    """文档管理器，用于管理 Markdown 文件和转换为 Word"""

    def __init__(self, base_dir: str = "./documents"):
        """
        初始化文档管理器

        Args:
            base_dir: 文档存储的基础目录
        """
        self.base_dir = Path(base_dir)
        self.base_dir.mkdir(parents=True, exist_ok=True)

    def _get_path(self, filename: str) -> Path:
        """获取文件的完整路径"""
        if not filename.endswith(".md"):
            filename += ".md"
        return self.base_dir / filename

    def create(self, filename: str, content: str = "") -> str:
        """
        创建新的 Markdown 文件

        Args:
            filename: 文件名（可不带 .md 后缀）
            content: 初始内容

        Returns:
            创建的文件路径
        """
        filepath = self._get_path(filename)
        if filepath.exists():
            raise FileExistsError(f"文件已存在: {filepath}")

        filepath.write_text(content, encoding="utf-8")
        return str(filepath)

    def read(self, filename: str) -> str:
        """
        读取 Markdown 文件内容

        Args:
            filename: 文件名

        Returns:
            文件内容
        """
        filepath = self._get_path(filename)
        if not filepath.exists():
            raise FileNotFoundError(f"文件不存在: {filepath}")

        return filepath.read_text(encoding="utf-8")

    def write(self, filename: str, content: str, append: bool = False) -> str:
        """
        写入 Markdown 文件

        Args:
            filename: 文件名
            content: 要写入的内容
            append: 是否追加模式，默认为覆盖

        Returns:
            文件路径
        """
        filepath = self._get_path(filename)

        if append and filepath.exists():
            existing = filepath.read_text(encoding="utf-8")
            content = existing + "\n" + content

        filepath.write_text(content, encoding="utf-8")
        return str(filepath)

    def delete(self, filename: str) -> bool:
        """
        删除 Markdown 文件

        Args:
            filename: 文件名

        Returns:
            是否删除成功
        """
        filepath = self._get_path(filename)
        if not filepath.exists():
            return False

        filepath.unlink()
        return True

    def list_files(self) -> list[str]:
        """
        列出所有 Markdown 文件

        Returns:
            文件名列表
        """
        return [f.name for f in self.base_dir.glob("*.md")]

    def exists(self, filename: str) -> bool:
        """检查文件是否存在"""
        return self._get_path(filename).exists()


class MarkdownToWordConverter:
    """Markdown 转 Word 转换器"""

    def __init__(self):
        self.doc = None

    def convert(self, markdown_content: str, output_path: str) -> str:
        """
        将 Markdown 内容转换为 Word 文档

        Args:
            markdown_content: Markdown 格式的文本
            output_path: 输出的 Word 文件路径

        Returns:
            生成的 Word 文件路径
        """
        self.doc = Document()

        lines = markdown_content.split("\n")
        i = 0

        while i < len(lines):
            line = lines[i]

            # 处理代码块
            if line.startswith("```"):
                code_lines = []
                lang = line[3:].strip()
                i += 1
                while i < len(lines) and not lines[i].startswith("```"):
                    code_lines.append(lines[i])
                    i += 1
                self._add_code_block("\n".join(code_lines), lang)
                i += 1
                continue

            # 处理标题
            if line.startswith("#"):
                level = len(line) - len(line.lstrip("#"))
                text = line.lstrip("#").strip()
                self._add_heading(text, level)

            # 处理无序列表
            elif line.strip().startswith("- ") or line.strip().startswith("* "):
                text = line.strip()[2:]
                self._add_list_item(text, ordered=False)

            # 处理有序列表
            elif re.match(r"^\d+\.\s", line.strip()):
                text = re.sub(r"^\d+\.\s", "", line.strip())
                self._add_list_item(text, ordered=True)

            # 处理引用
            elif line.startswith(">"):
                text = line.lstrip(">").strip()
                self._add_quote(text)

            # 处理水平线
            elif line.strip() in ["---", "***", "___"]:
                self._add_horizontal_line()

            # 处理普通段落
            elif line.strip():
                self._add_paragraph(line)

            i += 1

        # 确保输出目录存在
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        if not str(output_path).endswith(".docx"):
            output_path = Path(str(output_path) + ".docx")

        self.doc.save(str(output_path))
        return str(output_path)

    def _add_heading(self, text: str, level: int):
        """添加标题"""
        # Word 标题级别从 0 开始，Markdown 从 1 开始
        heading_level = min(level, 9)  # Word 最多支持 9 级标题
        self.doc.add_heading(text, level=heading_level)

    def _add_paragraph(self, text: str):
        """添加段落，处理内联格式"""
        para = self.doc.add_paragraph()
        self._parse_inline_formatting(para, text)

    def _parse_inline_formatting(self, paragraph, text: str):
        """解析并应用内联格式（粗体、斜体、代码等）"""
        # 处理粗体
        parts = re.split(r"\*\*(.+?)\*\*", text)
        for idx, part in enumerate(parts):
            if idx % 2 == 0:
                # 普通文本，继续处理斜体
                self._add_with_italic(paragraph, part)
            else:
                # 粗体文本
                run = paragraph.add_run(part)
                run.bold = True
                self._set_run_font(run)

    def _add_with_italic(self, paragraph, text: str):
        """处理斜体格式"""
        parts = re.split(r"\*(.+?)\*", text)
        for idx, part in enumerate(parts):
            if idx % 2 == 0:
                # 普通文本，继续处理行内代码
                self._add_with_code(paragraph, part)
            else:
                # 斜体文本
                run = paragraph.add_run(part)
                run.italic = True
                self._set_run_font(run)

    def _add_with_code(self, paragraph, text: str):
        """处理行内代码格式"""
        parts = re.split(r"`(.+?)`", text)
        for idx, part in enumerate(parts):
            if idx % 2 == 0:
                if part:
                    run = paragraph.add_run(part)
                    self._set_run_font(run)
            else:
                # 代码文本
                run = paragraph.add_run(part)
                run.font.name = "Consolas"
                run.font.size = Pt(10)

    def _set_run_font(self, run, font_name: str = "微软雅黑"):
        """设置 run 的字体，支持中文"""
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

    def _add_list_item(self, text: str, ordered: bool = False):
        """添加列表项"""
        style = "List Number" if ordered else "List Bullet"
        para = self.doc.add_paragraph(style=style)
        self._parse_inline_formatting(para, text)

    def _add_quote(self, text: str):
        """添加引用块"""
        para = self.doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.5)
        para.style = "Quote" if "Quote" in [s.name for s in self.doc.styles] else None
        self._parse_inline_formatting(para, text)

    def _add_code_block(self, code: str, language: str = ""):
        """添加代码块"""
        para = self.doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.3)
        run = para.add_run(code)
        run.font.name = "Consolas"
        run.font.size = Pt(9)

    def _add_horizontal_line(self):
        """添加水平线（用段落边框模拟）"""
        para = self.doc.add_paragraph()
        para.add_run("─" * 50)


def md_to_word(markdown_content: str, output_path: str) -> str:
    """
    便捷函数：将 Markdown 转换为 Word

    Args:
        markdown_content: Markdown 内容
        output_path: 输出路径

    Returns:
        生成的文件路径
    """
    converter = MarkdownToWordConverter()
    return converter.convert(markdown_content, output_path)


# 测试代码
if __name__ == "__main__":
    # 测试 DocumentManager
    dm = DocumentManager("./test_docs")

    # 创建文件
    dm.create("test", "# 测试文档\n\n这是一个测试。")
    print(f"文件列表: {dm.list_files()}")

    # 读取文件
    content = dm.read("test")
    print(f"文件内容:\n{content}")

    # 追加内容
    dm.write("test", "\n## 新章节\n\n追加的内容。", append=True)

    # 转换为 Word
    md_content = dm.read("test")
    md_to_word(md_content, "./test_docs/test.docx")
    print("Word 文档已生成")

    # 删除文件
    dm.delete("test")
    print(f"删除后文件列表: {dm.list_files()}")
